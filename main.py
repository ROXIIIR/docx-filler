import os
import json
import argparse
import re
from pathlib import Path
from docx import Document
from openai import OpenAI
from typing import List


# ------------------------------------------------------------
# CONFIG
# ------------------------------------------------------------
BLANK_RE = re.compile(r"[_\.]{3,}")  # detectează secvențe de ____ sau ....

client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))


# ------------------------------------------------------------
# 1. Încarcă JSON
# ------------------------------------------------------------
def load_json(json_path: Path) -> dict:
    with open(json_path, "r", encoding="utf-8") as f:
        return json.load(f)


# ------------------------------------------------------------
# 2. Extrage etichetele din DOCX
# ------------------------------------------------------------
def extract_labels_from_docx(template_path: Path) -> list[str]:
    """
    Extrage etichete pentru câmpuri din DOCX, inclusiv atunci când există
    mai multe blank-uri (_____ / ....) în același paragraf.

    Strategia:
    - pentru fiecare paragraf / celulă de tabel:
      - căutăm TOATE secvențele de blank (BLANK_RE.finditer)
      - pentru fiecare blank calculăm un label local:
        * textul dinainte de blank, fără alte blank-uri
        * dacă e gol și paragraful începe cu liniuțe, folosim last_label
    """
    doc = Document(str(template_path))
    labels = set()
    last_label = None

    def handle_text(text: str):
        nonlocal last_label
        if not text.strip():
            return

        # Găsim TOATE aparițiile de blank în același paragraf
        matches = list(BLANK_RE.finditer(text))
        if not matches:
            # Fără blank, poate fi doar etichetă simplă, o memorăm ca last_label
            last_label = text.strip()
            return

        # Dacă există blank-uri, procesăm fiecare în parte
        for i, m in enumerate(matches):
            before = text[:m.start()]
            # Scoatem alte blank-uri din before (dacă sunt)
            before_clean = BLANK_RE.sub("", before).strip()

            if before_clean:
                label = before_clean
            else:
                # dacă linia începe cu liniuțe și nu avem text înainte,
                # folosim ultima etichetă văzută (de ex. "Valabilitate ofertă")
                label = last_label

            if label:
                labels.add(label)
                last_label = label  # actualizăm last_label cu ultimul label folosit

        # Dacă după ultimul blank mai există text non-blank, îl putem memora ca last_label
        after = BLANK_RE.sub("", text[matches[-1].end():]).strip()
        if after:
            last_label = after

    # Parcurgem paragrafele și tabelele
    for p in doc.paragraphs:
        handle_text(p.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    handle_text(p.text)

    return sorted(labels)


# ------------------------------------------------------------
# 3. AGENTUL AI – LLM decide ce valoare merge în fiecare etichetă
# ------------------------------------------------------------
def llm_map_fields(labels: list[str], json_data: dict) -> dict:
    prompt = f"""
Ești un agent inteligent care completează formulare DOCX.

Primești:
1) O listă de etichete din document (labels)
2) Un JSON cu date (json_data)

Sarcina ta:
- Pentru fiecare etichetă, decide ce valoare din JSON trebuie inserată
- Dacă nu există un corespondent potrivit, întoarce "" (șir gol)
- Nu inventa câmpuri noi
- Returnează STRICT un obiect JSON Python dict

Labels:
{json.dumps(labels, ensure_ascii=False, indent=2)}

JSON data:
{json.dumps(json_data, ensure_ascii=False, indent=2)}
    """

    response = client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[
            {"role": "system", "content": "Completezi DOCX-uri pe baza JSON-ului."},
            {"role": "user", "content": prompt},
        ],
        response_format={"type": "json_object"},
    )

    return json.loads(response.choices[0].message.content)


# ------------------------------------------------------------
# 4. Completează DOCX-ul pe baza mapping-ului dat de LLM
# ------------------------------------------------------------
def fill_docx_with_mapping(template_path: Path, mapping: dict, output_path: Path):
    """
    Completează DOCX-ul folosind mapping-ul LLM-ului.
    Suportă mai multe blank-uri (_____ / ....) în același paragraf.
    """
    doc = Document(str(template_path))
    last_label = None

    def fill_paragraph(p):
        nonlocal last_label
        text = p.text
        if not text.strip():
            return

        matches = list(BLANK_RE.finditer(text))
        if not matches:
            # paragraful poate conține doar o etichetă
            last_label = text.strip()
            return

        # Vom reconstrui textul paragrafului, înlocuind fiecare blank
        new_text = ""
        cursor = 0

        for i, m in enumerate(matches):
            before = text[cursor:m.start()]
            before_clean = BLANK_RE.sub("", text[:m.start()]).strip()

            if before_clean:
                label = before_clean
            else:
                label = last_label

            # adăugăm textul de dinainte de blank, nemodificat
            new_text += text[cursor:m.start()]

            value = ""
            if label and label in mapping:
                value = mapping[label]

            # în locul blank-ului punem valoarea (sau blank gol dacă nu avem)
            new_text += value or ""

            # actualizăm last_label și cursorul
            if label:
                last_label = label
            cursor = m.end()

        # adăugăm restul de text după ultimul blank
        new_text += text[cursor:]

        # suprascriem run-urile paragrafului cu noul text (simplu, dar sigur)
        if p.runs:
            p.runs[0].text = new_text
            for r in p.runs[1:]:
                r.text = ""
        else:
            p.add_run(new_text)

    # aplicăm pentru paragrafe
    for p in doc.paragraphs:
        fill_paragraph(p)

    # și pentru tabele
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    fill_paragraph(p)

    doc.save(str(output_path))
    print(f"[DONE] Saved filled DOCX → {output_path}")


# ------------------------------------------------------------
# 5. MAIN
# ------------------------------------------------------------
def main():
    parser = argparse.ArgumentParser(description="Agentic AI DOCX filler")
    parser.add_argument("--template", required=True, help="Template DOCX")
    parser.add_argument("--data", required=True, help="JSON data file")
    parser.add_argument("--output", required=False, help="Output DOCX file")

    args = parser.parse_args()

    template_path = Path(args.template)
    json_path = Path(args.data)

    output_path = Path(args.output) if args.output else template_path.with_name(
        template_path.stem + ".llm_filled.docx"
    )

    print("[1] Loading JSON...")
    json_data = load_json(json_path)

    print("[2] Extracting labels from DOCX...")
    labels = extract_labels_from_docx(template_path)

    print("[3] Asking LLM for mapping...")
    mapping = llm_map_fields(labels, json_data)

    print("[4] Filling DOCX with mapping...")
    fill_docx_with_mapping(template_path, mapping, output_path)


if __name__ == "__main__":
    main()
